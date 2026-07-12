import uuid
from io import BytesIO
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from database.db import SessionLocal
from models.complaint import Complaint
from models.fir_drafts import FirDraft


router = APIRouter(prefix="/api/fir-drafts", tags=["fir-drafts"])


class FirDraftCreate(BaseModel):
    complaint_id: str
    crime_category: Optional[str] = None
    summary: Optional[str] = None
    draft_content: Optional[Dict[str, Any]] = None
    status: Optional[str] = "draft"
    officer_notes: Optional[str] = None
    approved_by: Optional[str] = None
    approved_at: Optional[datetime] = None


class FirDraftUpdate(BaseModel):
    complaint_id: Optional[str] = None
    crime_category: Optional[str] = None
    summary: Optional[str] = None
    draft_content: Optional[Dict[str, Any]] = None
    status: Optional[str] = None
    officer_notes: Optional[str] = None
    approved_by: Optional[str] = None
    approved_at: Optional[datetime] = None


class FirDraftSummary(BaseModel):
    id: str
    complaint_id: str
    crime_category: Optional[str] = None
    summary: Optional[str] = None
    draft_content: Optional[Dict[str, Any]] = None
    status: Optional[str] = None
    officer_notes: Optional[str] = None
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None


class FirDraftCreateData(BaseModel):
    id: str


class FirDraftCreateResponse(BaseModel):
    success: bool
    message: str
    data: FirDraftCreateData


class FirDraftListResponse(BaseModel):
    fir_drafts: List[FirDraftSummary]


class FirDraftDetailResponse(FirDraftSummary):
    pass


class FirDraftUpdateResponse(BaseModel):
    success: bool
    message: str
    data: FirDraftCreateData


class FirDraftDeleteResponse(BaseModel):
    success: bool
    message: str


@router.post(
    "",
    response_model=FirDraftCreateResponse,
    summary="Create an FIR draft",
    description="Create a new FIR draft record for a complaint.",
    tags=["fir-drafts"],
)
def create_fir_draft(payload: FirDraftCreate) -> Dict[str, Any]:
    db = SessionLocal()
    try:
        fir_draft = FirDraft(
            complaint_id=uuid.UUID(payload.complaint_id),
            crime_category=payload.crime_category,
            summary=payload.summary,
            draft_content=payload.draft_content if payload.draft_content is not None else {},
            status=payload.status if payload.status is not None else "draft",
            officer_notes=payload.officer_notes,
            approved_by=payload.approved_by,
            approved_at=payload.approved_at
        )
        
        db.add(fir_draft)
        db.commit()
        db.refresh(fir_draft)
        
        return {
            "success": True,
            "message": "FIR draft created successfully",
            "data": {
                "id": str(fir_draft.id)
            }
        }
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID format")
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to create FIR draft: {str(e)}")
    finally:
        db.close()


@router.get(
    "",
    response_model=FirDraftListResponse,
    summary="List FIR drafts",
    description="Retrieve all FIR draft records.",
    tags=["fir-drafts"],
)
def get_all_fir_drafts() -> Dict[str, Any]:
    db = SessionLocal()
    try:
        fir_drafts = db.query(FirDraft).all()
        return {
            "fir_drafts": [
                {
                    "id": str(fd.id),
                    "complaint_id": str(fd.complaint_id),
                    "crime_category": fd.crime_category,
                    "summary": fd.summary,
                    "draft_content": fd.draft_content,
                    "status": fd.status,
                    "officer_notes": fd.officer_notes,
                    "approved_by": fd.approved_by,
                    "approved_at": fd.approved_at.isoformat() if fd.approved_at else None,
                    "created_at": fd.created_at.isoformat() if fd.created_at else None,
                    "updated_at": fd.updated_at.isoformat() if fd.updated_at else None
                }
                for fd in fir_drafts
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve FIR drafts: {str(e)}")
    finally:
        db.close()


@router.get(
    "/{id}",
    response_model=FirDraftDetailResponse,
    summary="Get FIR draft",
    description="Retrieve a single FIR draft record by its identifier.",
    tags=["fir-drafts"],
)
def get_fir_draft(id: str) -> Dict[str, Any]:
    db = SessionLocal()
    try:
        fir_draft = db.query(FirDraft).filter(FirDraft.id == uuid.UUID(id)).first()
        if not fir_draft:
            raise HTTPException(status_code=404, detail="FIR draft not found")
        
        return {
            "id": str(fir_draft.id),
            "complaint_id": str(fir_draft.complaint_id),
            "crime_category": fir_draft.crime_category,
            "summary": fir_draft.summary,
            "draft_content": fir_draft.draft_content,
            "status": fir_draft.status,
            "officer_notes": fir_draft.officer_notes,
            "approved_by": fir_draft.approved_by,
            "approved_at": fir_draft.approved_at.isoformat() if fir_draft.approved_at else None,
            "created_at": fir_draft.created_at.isoformat() if fir_draft.created_at else None,
            "updated_at": fir_draft.updated_at.isoformat() if fir_draft.updated_at else None
        }
    except HTTPException:
        raise
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid id format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve FIR draft: {str(e)}")
    finally:
        db.close()


@router.put(
    "/{id}",
    response_model=FirDraftUpdateResponse,
    summary="Update an FIR draft",
    description="Partially update an existing FIR draft record.",
    tags=["fir-drafts"],
)
def update_fir_draft(id: str, payload: FirDraftUpdate) -> Dict[str, Any]:
    db = SessionLocal()
    try:
        fir_draft = db.query(FirDraft).filter(FirDraft.id == uuid.UUID(id)).first()
        if not fir_draft:
            raise HTTPException(status_code=404, detail="FIR draft not found")
        
        # Update only provided fields
        if payload.complaint_id is not None:
            fir_draft.complaint_id = uuid.UUID(payload.complaint_id)
        if payload.crime_category is not None:
            fir_draft.crime_category = payload.crime_category
        if payload.summary is not None:
            fir_draft.summary = payload.summary
        if payload.draft_content is not None:
            fir_draft.draft_content = payload.draft_content
        if payload.status is not None:
            fir_draft.status = payload.status
        if payload.officer_notes is not None:
            fir_draft.officer_notes = payload.officer_notes
        if payload.approved_by is not None:
            fir_draft.approved_by = payload.approved_by
        if payload.approved_at is not None:
            fir_draft.approved_at = payload.approved_at
        
        db.commit()
        db.refresh(fir_draft)
        
        return {
            "success": True,
            "message": "FIR draft updated successfully",
            "data": {
                "id": str(fir_draft.id)
            }
        }
    except HTTPException:
        raise
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID format")
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update FIR draft: {str(e)}")
    finally:
        db.close()


@router.delete(
    "/{id}",
    response_model=FirDraftDeleteResponse,
    summary="Delete an FIR draft",
    description="Remove an FIR draft record from the system.",
    tags=["fir-drafts"],
)
def delete_fir_draft(id: str) -> Dict[str, Any]:
    db = SessionLocal()
    try:
        fir_draft = db.query(FirDraft).filter(FirDraft.id == uuid.UUID(id)).first()
        if not fir_draft:
            raise HTTPException(status_code=404, detail="FIR draft not found")
        
        db.delete(fir_draft)
        db.commit()
        
        return {
            "success": True,
            "message": "FIR draft deleted successfully"
        }
    except HTTPException:
        raise
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid id format")
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete FIR draft: {str(e)}")
    finally:
        db.close()


@router.get(
    "/{id}/download",
    summary="Download FIR draft as DOCX",
    description="Generate a DOCX document for an FIR draft and download it.",
    response_class=StreamingResponse,
    tags=["fir-drafts"],
)
def download_fir_draft(id: str):
    db = SessionLocal()
    try:
        fir_draft = db.query(FirDraft).filter(FirDraft.id == uuid.UUID(id)).first()
        if not fir_draft:
            raise HTTPException(status_code=404, detail="FIR draft not found")

        complaint = db.query(Complaint).filter(Complaint.complaint_id == str(fir_draft.complaint_id)).first()
        document = Document()

        document.add_heading(f"FIR Draft for Complaint {fir_draft.complaint_id}", level=1)
        document.add_paragraph(f"FIR Draft ID: {fir_draft.id}")
        document.add_paragraph(f"Complaint ID: {fir_draft.complaint_id}")
        document.add_paragraph(f"Crime category: {fir_draft.crime_category or 'Not specified'}")
        document.add_paragraph(f"Status: {fir_draft.status}")
        document.add_paragraph("")

        if complaint:
            document.add_heading("Complaint details", level=2)
            document.add_paragraph(f"Complainant: {complaint.complainant_name or 'Not provided'}")
            document.add_paragraph(f"Phone: {complaint.phone or 'Not provided'}")
            document.add_paragraph(f"Email: {complaint.email or 'Not provided'}")
            document.add_paragraph(f"Location: {complaint.location or 'Not provided'}")
            document.add_paragraph(f"Description: {complaint.description or 'Not provided'}")
            document.add_paragraph("")

        if fir_draft.summary:
            document.add_heading("Case summary", level=2)
            document.add_paragraph(fir_draft.summary)
            document.add_paragraph("")

        draft_content = fir_draft.draft_content or {}
        selected_sections = draft_content.get("selected_sections", [])
        if selected_sections:
            document.add_heading("Selected legal sections", level=2)
            for section in selected_sections:
                document.add_paragraph(
                    f"{section.get('act_code', '')} Sec {section.get('section_number', '')} — {section.get('title', '')}",
                    style="List Bullet",
                )
                if section.get("reason"):
                    document.add_paragraph(f"Reason: {section['reason']}")
                cross_references = section.get("cross_references", [])
                if cross_references:
                    document.add_paragraph("Cross references:")
                    for xref in cross_references:
                        document.add_paragraph(
                            f"- {xref.get('act')} Sec {xref.get('section')} {xref.get('subject', '')}",
                        )
            document.add_paragraph("")

        selected_judgments = draft_content.get("selected_judgments", [])
        if selected_judgments:
            document.add_heading("Selected landmark judgments", level=2)
            for judgment in selected_judgments:
                document.add_paragraph(
                    f"{judgment.get('case_title', '')} ({judgment.get('court', '')}, {judgment.get('case_date', '')})",
                    style="List Bullet",
                )
                if judgment.get("ipc_sections"):
                    document.add_paragraph(f"IPC Sections: {judgment['ipc_sections']}")
                if judgment.get("judgment_reason"):
                    document.add_paragraph(f"Reason: {judgment['judgment_reason']}")
                if judgment.get("summary"):
                    document.add_paragraph(f"Summary: {judgment['summary']}")
            document.add_paragraph("")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"fir_draft_{id}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            },
        )
    except HTTPException:
        raise
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid id format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX document: {str(e)}")
    finally:
        db.close()
