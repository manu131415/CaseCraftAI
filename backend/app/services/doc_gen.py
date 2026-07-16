import os
import copy
import uuid
import datetime
from docxtpl import DocxTemplate
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv

# Load database configuration
ENV_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
load_dotenv(dotenv_path=ENV_PATH)
DATABASE_URL = os.getenv("DATABASE_URL")

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated_docs")

def get_db_conn():
    return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)

def duplicate_row(table, template_row):
    """Helper to duplicate a table row and preserve style."""
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    from docx.table import _Row
    return _Row(new_tr, table)

def replace_text_in_paragraph(p, replacements):
    """Replaces text in paragraph while preserving styles as much as possible."""
    combined_text = p.text
    modified = False
    for key, val in replacements.items():
        if key in combined_text:
            combined_text = combined_text.replace(key, str(val) if val is not None else "")
            modified = True
    if modified:
        p.text = combined_text

def fill_docx_template(template_name, output_filename, replacements, evidences=None):
    """Loads a template using DocxTemplate, populates the fields, and saves it."""
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Load using DocxTemplate as requested
    tpl = DocxTemplate(template_path)
    tpl.init_docx()
    
    # Process regular paragraphs in the DocxTemplate object
    for p in tpl.paragraphs:
        replace_text_in_paragraph(p, replacements)
        
    # Process tables in the DocxTemplate object
    for table in tpl.tables:
        # Check if this table has the evidence placeholder row
        has_evidence_placeholder = False
        template_row = None
        for idx, row in enumerate(table.rows):
            if any("[e.g., Apple iPhone 15 Pro" in cell.text for cell in row.cells):
                has_evidence_placeholder = True
                template_row = row
                break
        
        if has_evidence_placeholder and evidences is not None:
            # Handle evidence list dynamically
            if not evidences:
                for cell in template_row.cells:
                    cell.text = cell.text.replace("[e.g., Apple iPhone 15 Pro, IMEI: 35...]", "No Seized Material")
                    cell.text = cell.text.replace("1", "0")
                    cell.text = cell.text.replace("Switched Off, Packaged in Plastic Bag 'A'", "N/A")
            else:
                for idx, ev in enumerate(evidences):
                    if idx == 0:
                        row_to_fill = template_row
                    else:
                        row_to_fill = duplicate_row(table, template_row)
                    
                    row_to_fill.cells[0].text = str(idx + 1)
                    row_to_fill.cells[1].text = f"{ev.get('evidence_type', '')} - {ev.get('description', '')}\nSerial No: {ev.get('serial_number', 'N/A')}"
                    row_to_fill.cells[2].text = str(ev.get('quantity', 1))
                    row_to_fill.cells[3].text = f"Condition: {ev.get('item_condition', 'N/A')}\nSeal No: {ev.get('seal_number', 'N/A')}"
        else:
            # Process table cells normally
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, replacements)
                        
    # Process sections (headers/footers) in the DocxTemplate object
    for section in tpl.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    replace_text_in_paragraph(p, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    replace_text_in_paragraph(p, replacements)
                    
    # Call render on the template to complete docxtpl workflow
    tpl.render(replacements)
    
    # Save the generated document
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    tpl.save(output_path)
    return output_path

def log_generated_document(conn, case_id, doc_type, title, file_path, metadata_dict):
    """Inserts a row into the documents table."""
    cursor = conn.cursor()
    document_id = str(uuid.uuid4())
    generated_at = datetime.datetime.now(datetime.timezone.utc)
    
    # Query case's assigned officer
    cursor.execute("SELECT assigned_officer_id FROM cases WHERE case_id = %s;", (case_id,))
    row = cursor.fetchone()
    generated_by = row['assigned_officer_id'] if row else "SYSTEM"
    
    import json
    metadata_json = json.dumps(metadata_dict, default=str)
    
    cursor.execute("""
        INSERT INTO documents (
            document_id, case_id, document_type, title, file_path, status, generated_by, generated_at, version, document_metadata
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
    """, (
        document_id,
        case_id,
        doc_type,
        title,
        file_path,
        "Generated",
        generated_by,
        generated_at,
        "1.0",
        metadata_json
    ))
    conn.commit()
    cursor.close()
    return document_id

def format_date(d):
    if not d:
        return ""
    if isinstance(d, (datetime.date, datetime.datetime)):
        return d.strftime("%d/%m/%Y")
    return str(d)

def format_datetime(dt):
    if not dt:
        return ""
    if isinstance(dt, datetime.datetime):
        return dt.strftime("%d/%m/%Y at %H:%M")
    return str(dt)

def generate_document_by_type(case_id, doc_type, accused_id=None):
    """Query data from DB, construct replacements, fill template using DocxTemplate, and log it."""
    conn = get_db_conn()
    cursor = conn.cursor()
    
    try:
        # 1. Fetch Case
        cursor.execute("SELECT * FROM cases WHERE case_id = %s;", (case_id,))
        case = cursor.fetchone()
        if not case:
            raise ValueError(f"Case with ID {case_id} not found.")
            
        # 2. Fetch Officer
        officer = None
        if case.get('assigned_officer_id'):
            cursor.execute("SELECT * FROM officers WHERE officer_id = %s;", (case['assigned_officer_id'],))
            officer = cursor.fetchone()
            
        # 3. Fetch Accused
        accused_list = []
        if accused_id:
            cursor.execute("SELECT * FROM accused WHERE accused_id = %s AND case_id = %s;", (accused_id, case_id))
            accused_list = cursor.fetchall()
        else:
            cursor.execute("SELECT * FROM accused WHERE case_id = %s ORDER BY created_at ASC;", (case_id,))
            accused_list = cursor.fetchall()
            
        accused = accused_list[0] if accused_list else None
        
        # 4. Fetch Evidence
        cursor.execute("SELECT * FROM evidences WHERE case_id = %s;", (case_id,))
        evidences = cursor.fetchall()
        
        # 5. Fetch Witnesses
        cursor.execute("SELECT * FROM witnesses WHERE case_id = %s;", (case_id,))
        witnesses = cursor.fetchall()
        
        # 6. Fetch Victims
        cursor.execute("SELECT * FROM victims WHERE case_id = %s;", (case_id,))
        victims = cursor.fetchall()
        
        # 7. Fetch Legal Sections
        cursor.execute("""
            SELECT ls.section_number, ls.act_code 
            FROM case_legal_sections cls
            JOIN legal_sections ls ON cls.legal_section_id = ls.id
            WHERE cls.case_id = %s;
        """, (case_id,))
        sections = cursor.fetchall()
        sections_str = ", ".join([f"Sec {s['section_number']} {s['act_code']}" for s in sections])
        if not sections_str:
            sections_str = "Sections 318, 319 BNS" # Default fallback
            
        # Common variables
        current_year = str(datetime.datetime.now().year)
        today_str = datetime.date.today().strftime("%d/%m/%Y")
        
        officer_name = officer['name'] if officer else "Rajesh Patel"
        officer_rank = officer['rank'] if officer else "Inspector"
        officer_station = officer['station'] if officer else "Cyber Crime Police Station Ahmedabad"
        officer_badge = officer['badge_number'] if officer else "GJ-4521"
        officer_full = f"{officer_rank} {officer_name}"
        
        accused_name = accused['full_name'] if accused else "N/A"
        accused_alias = accused['alias'] if accused else ""
        accused_full_name = f"{accused_name} @ {accused_alias}" if accused_alias else accused_name
        accused_father = accused['father_name'] if accused else "N/A"
        accused_age = str(accused['age']) if accused and accused['age'] else "N/A"
        accused_gender = accused['gender'] if accused else "Male"
        accused_present = accused['present_address'] if accused and accused['present_address'] else "N/A"
        accused_perm = accused['permanent_address'] if accused and accused['permanent_address'] else "N/A"
        accused_addr = f"Present: {accused_present}; Permanent: {accused_perm}"
        accused_arrest = format_datetime(accused['arrest_datetime']) if accused else "N/A"
        accused_status = accused['custody_status'] if accused and accused['custody_status'] else "In Judicial Custody"
        
        fir_no = case['fir_no'] if case['fir_no'] else "N/A"
        fir_year = str(case['fir_year']) if case['fir_year'] else current_year
        fir_date = format_date(case['fir_date']) if case['fir_date'] else today_str
        police_station = case['police_station'] if case['police_station'] else officer_station
        district = case['district'].capitalize() if case['district'] else "Ahmedabad City"
        
        replacements = {}
        output_file = ""
        title = ""
        
        if doc_type == "purvani_chargesheet":
            # 1. Purvani Chargesheet
            cursor.execute("""
                SELECT * FROM cases 
                WHERE case_id = %s AND supplementary_chargesheet_no IS NOT NULL;
            """, (case_id,))
            supp_case = cursor.fetchone()
            
            supp_no = supp_case['supplementary_chargesheet_no'] if supp_case else "SUP-01"
            supp_reason = supp_case['supplementary_reason'] if supp_case else "Further investigation in cyber financial networks"
            orig_cs_no = case['original_chargesheet_no'] if case['original_chargesheet_no'] else "CS-101"
            orig_cs_date = format_date(case['original_chargesheet_date']) if case['original_chargesheet_date'] else today_str
            
            # Witness statements addendum
            witness_statements = []
            for w in witnesses:
                witness_statements.append(f"Witness {w['full_name']}: {w['statement']}")
            witness_str = "\n".join(witness_statements) if witness_statements else "Statements of additional panch witnesses recorded."
            
            # Additional Seized Material
            ev_list = []
            for ev in evidences:
                ev_list.append(f"{ev['evidence_type']} - {ev['description']} (Serial: {ev['serial_number']})")
            ev_str = "\n".join(ev_list) if ev_list else "Digital trail logs recovered from investment server."
            
            replacements = {
                "[___]/[Year]": f"{supp_no}/{current_year}",
                "[e.g., Naroda / Cyber Crime]": police_station,
                "[____]/[Year]": f"{fir_no}/{fir_year}",
                "[DD/MM/YYYY]": today_str,
                "Date of Original FIR: [DD/MM/YYYY]": f"Date of Original FIR: {fir_date}",
                "[____] / [DD/MM/YYYY]": f"{orig_cs_no} / {orig_cs_date}",
                "[e.g., Forensic Labs Report / Arrest of Absconding Accused]": supp_reason,
                "[Rank & Name]": officer_full,
                "[___________________]": accused_full_name,
                "[_______________________]": accused_father,
                "[___]": accused_age,
                "[Male/Female/Other]": accused_gender,
                "[_______________________________________]": accused_addr,
                "[DD/MM/YYYY at HH:MM]": accused_arrest,
                "[In Judicial Custody / Out on Bail]": accused_status,
                "[Brief of new witness statements]": witness_str,
                "[List and descriptions]": ev_str,
                "[e.g., 318, 319 of BNS]": sections_str
            }
            output_file = f"purvani_chargesheet_{case_id}.docx"
            title = f"Purvani Chargesheet - FIR {fir_no}/{fir_year}"
            fill_docx_template("purvani_chargesheet.docx", output_file, replacements)
            
        elif doc_type == "medical_treatment_letter":
            # 2. Medical Treatment Letter
            cursor.execute("SELECT * FROM medical_reports WHERE case_id = %s;", (case_id,))
            med_rep = cursor.fetchone()
            
            hospital = med_rep['hospital_name'] if med_rep else "Civil Hospital Asarwa"
            doctor = med_rep['doctor_name'] if med_rep else "Dr. A. K. Sharma"
            injuries = med_rep['visible_injuries'] if med_rep else "No visible external injuries, complaints of stress"
            injury_type = med_rep['injury_type'] if med_rep else "Simple"
            
            # Determine patient
            patient_name = accused_full_name
            patient_age_gender = f"{accused_age} / {accused_gender}"
            legal_status = "Accused under Custody"
            
            if med_rep and med_rep['victim_id']:
                cursor.execute("SELECT * FROM victims WHERE victim_id = %s;", (med_rep['victim_id'],))
                vic = cursor.fetchone()
                if vic:
                    patient_name = vic['full_name']
                    patient_age_gender = f"{vic['age']} / {vic['gender']}"
                    legal_status = "Injured Victim"
            
            replacements = {
                "[POLICE STATION NAME]": police_station.upper(),
                "[_________]/[Year]": f"MED-{fir_no}/{current_year}",
                "[DD/MM/YYYY]": today_str,
                "[e.g., Civil Hospital, Asarwa / LG Hospital]": hospital,
                "[___________________________]": patient_name,
                "[____] / [M/F]": patient_age_gender,
                "[Accused under Custody / Complainant / Injured Victim]": legal_status,
                "[_______]/[Year]": f"{fir_no}/{fir_year}",
                "[_________] BNS": sections_str,
                "[Briefly list visible trauma or physical distress]": injuries,
                "[Name & Buckle No. of PC/HC]": f"PC Ramesh Kumar, Buckle No. {officer_badge}",
                "[Seal of Ahmedabad Police Station]": police_station
            }
            output_file = f"medical_letter_{case_id}.docx"
            title = f"Medical Treatment Letter - {patient_name}"
            fill_docx_template("medical_treatment_letter.docx", output_file, replacements)
            
        elif doc_type == "remand_request_letter":
            # 3. Remand Request Letter
            cursor.execute("SELECT * FROM remand_details WHERE case_id = %s;", (case_id,))
            remand = cursor.fetchone()
            
            remand_days = str(remand['remand_days']) if remand else "7"
            grounds = remand['grounds'] if remand else "Interrogation regarding cyber transactions and decryption of hidden drives."
            
            court_no = case['court_no'] if case['court_no'] else "Court No. 3"
            
            # Calculate mandatory 24 hours
            expire_str = "N/A"
            if accused and accused['arrest_datetime']:
                expire_dt = accused['arrest_datetime'] + datetime.timedelta(hours=24)
                expire_str = format_datetime(expire_dt)
                
            replacements = {
                "[____]": court_no,
                "[_____]/[Year]": f"REM-{fir_no}/{current_year}",
                "State of Gujarat (Through [Name] Police Station)": f"State of Gujarat (Through {police_station})",
                "[Name] Police Station": police_station,
                "[Accused Person's Full Name]": accused_full_name,
                "[___] DAYS": f"{remand_days} DAYS",
                "[DD/MM/YYYY]": today_str,
                "[DD/MM/YYYY] at [HH:MM]": accused_arrest,
                "arrested on [DD/MM/YYYY] at [HH:MM]": f"arrested on {accused_arrest}",
                "expire on [DD/MM/YYYY] at [HH:MM]": f"expire on {expire_str}",
                "[____]/[Year]": f"{fir_no}/{fir_year}",
                "[___________]": sections_str,
                "[Location]": district,
                "[remand_reasons]": grounds,
                "[investigation_status]": "Active investigation under progress.",
                "[pending_investigation]": "Recovery of the source codes and electronic devices.",
                "remanded to Police Custody for a period of [___] days": f"remanded to Police Custody for a period of {remand_days} days",
                "[Rank & Name]": officer_full
            }
            output_file = f"remand_request_{case_id}.docx"
            title = f"Remand Request - {accused_name}"
            fill_docx_template("remand_request_letter.docx", output_file, replacements)
            
        elif doc_type == "seizure_receipt":
            # 4. Seizure Receipt
            seizure_loc = "Ahmedabad City"
            seizure_dt = today_str
            seized_from = accused_full_name
            if evidences:
                ev = evidences[0]
                seizure_loc = ev['seizure_location'] if ev['seizure_location'] else seizure_loc
                seizure_dt = format_datetime(ev['seizure_datetime']) if ev['seizure_datetime'] else seizure_dt
                seized_from = ev['seized_from'] if ev['seized_from'] else seized_from
                
            replacements = {
                "[_______]/[Year]": f"{fir_no}/{fir_year}",
                "[__________________]": police_station,
                "[DD/MM/YYYY] at [HH:MM]": seizure_dt,
                "[Full Address / Spot Details]": seizure_loc,
                "[___________________________]": seized_from,
                "[_______________________________________________________]": accused_present,
                "[Rank & Name]": officer_full,
                "[Name]": officer_name,
                "Name: [___________________________]": f"Name: {seized_from}",
                "Address: [_______________________________________________________]": f"Address: {accused_present}"
            }
            output_file = f"seizure_receipt_{case_id}.docx"
            title = f"Seizure Receipt - FIR {fir_no}/{fir_year}"
            fill_docx_template("seizure_receipt.docx", output_file, replacements, evidences=evidences)
            
        elif doc_type == "court_custody_letter":
            # 5. Court Custody Letter
            cursor.execute("SELECT * FROM court_custody WHERE case_id = %s;", (case_id,))
            custody = cursor.fetchone()
            
            prison = custody['prison_name'] if custody else "Sabarmati Central Jail"
            commitment_from = format_date(custody['commitment_from']) if custody else today_str
            commitment_to = format_date(custody['commitment_to']) if custody else format_date(datetime.date.today() + datetime.timedelta(days=14))
            court_order_no = custody['court_order_number'] if custody else "MET-ORD-452"
            
            court_no = case['court_no'] if case['court_no'] else "Court No. 3"
            
            replacements = {
                "[POLICE STATION]": police_station.upper(),
                "[Accused Name]": accused_full_name,
                "[__]": court_no,
                "[DD/MM/YYYY]": today_str,
                "[Name, Age, Identification Mark]": f"{accused_full_name}, Age: {accused_age}, Marks: {accused['identification_marks'] if accused and accused['identification_marks'] else 'None'}",
                "[______]/[Year]": f"{fir_no}/{fir_year}",
                "[_________]": police_station,
                "From [DD/MM/YYYY] to [DD/MM/YYYY]": f"From {commitment_from} to {commitment_to}",
                "dated [DD/MM/YYYY]": f"dated {commitment_from}"
            }
            output_file = f"court_custody_{case_id}.docx"
            title = f"Court Custody Letter - {accused_name}"
            fill_docx_template("court_custody_request.docx", output_file, replacements)
            
        elif doc_type == "accused_panchanama":
            # 6. Accused Panchanama
            panch1 = "Vijay Sharma, Age: 34, Occ: Business, Addr: 12, GIDC, Naroda, Aadhaar: 4521-7896-1234"
            panch2 = "Suresh Patel, Age: 42, Occ: Service, Addr: 45, Sterling Appt, Vastrapur, Aadhaar: 7894-1235-4569"
            if len(witnesses) > 0:
                w = witnesses[0]
                panch1 = f"{w['full_name']}, Address: {w['address']}"
            if len(witnesses) > 1:
                w = witnesses[1]
                panch2 = f"{w['full_name']}, Address: {w['address']}"
                
            ev_list = []
            for ev in evidences:
                ev_list.append(f"{ev['quantity']} unit(s) of {ev['evidence_type']} ({ev['description']})")
            ev_str = ", ".join(ev_list) if ev_list else "one smartphone and bank records"
            
            replacements = {
                "[DD/MM/YYYY]": today_str,
                "[HH:MM]": "10:00 AM",
                "Concluded at: [HH:MM]": "Concluded at: 01:00 PM",
                "[Exact Spot Description, Ahmedabad]": district,
                "[_______] Police Station": police_station,
                "Panch No 1: [Name, Age, Occupation, Full Residential Address, Aadhaar No]": f"Panch No 1: {panch1}",
                "Panch No 2: [Name, Age, Occupation, Full Residential Address, Aadhaar No]": f"Panch No 2: {panch2}",
                "[Accused Name]": accused_full_name,
                "[found_articles]": ev_str
            }
            output_file = f"accused_panchanama_{case_id}.docx"
            title = f"Accused Panchanama - {accused_name}"
            fill_docx_template("accused_panchanama.docx", output_file, replacements)
            
        elif doc_type == "accused_face_identification_form":
            # 7. Accused Face Identification Form
            face_shape = accused['face_shape'] if accused and accused['face_shape'] else "Oval"
            complexion = accused['complexion'] if accused and accused['complexion'] else "Wheatish"
            eye_color = accused['eye_color'] if accused and accused['eye_color'] else "Black"
            eye_structure = accused['eye_structure'] if accused and accused['eye_structure'] else "Standard"
            hair_type = accused['hair_type'] if accused and accused['hair_type'] else "Straight"
            hair_color = accused['hair_color'] if accused and accused['hair_color'] else "Black"
            capture_date = format_date(accused['capture_date']) if accused and accused['capture_date'] else today_str
            
            marks = accused['identification_marks'].split(";") if accused and accused['identification_marks'] else []
            mark_a = marks[0] if len(marks) > 0 else "Cut mark on forehead"
            mark_b = marks[1] if len(marks) > 1 else "Tattoo on left forearm"
            
            replacements = {
                "[__________]": accused_id if accused_id else "ACC-01",
                "[__________]/[Year]": f"{fir_no}/{fir_year}",
                "[__________________]": police_station,
                "[DD/MM/YYYY]": capture_date,
                "[_______________________]": accused_name,
                "[_____________________]": accused_alias if accused_alias else "N/A",
                "[_____] / [M/F/Other]": f"{accused_age} / {accused_gender}",
                "[Oval / Round / Square / Long]": face_shape,
                "[Fair / Wheatish / Dark]": complexion,
                "[Black / Brown / Grey] | [Standard / Deep-set]": f"{eye_color} | {eye_structure}",
                "[Straight / Curly / Bald / Grey / Black]": f"{hair_type} / {hair_color}",
                "[e.g., Deep linear scar on right jawline]": mark_a,
                "[e.g., Tattoo of alphabets 'OM' on left wrist]": mark_b,
                "FR_FRONT_[ID].jpg": f"FR_FRONT_{accused_id}.jpg",
                "FR_RPROF_[ID].jpg": f"FR_RPROF_{accused_id}.jpg",
                "FR_LPROF_[ID].jpg": f"FR_LPROF_{accused_id}.jpg",
                "[Name & Designation of Officer]": officer_full
            }
            output_file = f"face_id_{case_id}.docx"
            title = f"Face Identification Form - {accused_name}"
            fill_docx_template("accused_face_identification_form.docx", output_file, replacements)
            
        else:
            raise ValueError(f"Unknown document type: {doc_type}")
            
        # Log to documents table
        rel_output_path = os.path.join("app", "services", "generated_docs", output_file)
        doc_id = log_generated_document(conn, case_id, doc_type, title, rel_output_path, replacements)
        
        return {
            "document_id": doc_id,
            "case_id": case_id,
            "document_type": doc_type,
            "title": title,
            "file_path": rel_output_path,
            "status": "Generated",
            "generated_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
    finally:
        cursor.close()
        conn.close()
