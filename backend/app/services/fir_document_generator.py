"""Renders a `FirDraft` row into the same First Information Report layout
as the blank FIR template (District/PS block, Act & Sections, Occurrence,
Place, Complainant, Accused table, ... through Action Taken and
signatures), populated with whatever data is available.

Any field the draft/complaint doesn't have is left as a blank underline,
exactly as it appears on the blank template — this is meant to be printed
and completed by hand/officer sign-off where the system doesn't hold the
data (e.g. accused address, property particulars).
"""
from __future__ import annotations

import io
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HEADER_FILL = "E8E8E8"
LABEL_SIZE = Pt(9)
SUBLABEL_SIZE = Pt(7.5)


# ---------------------------------------------------------------------------
# small helpers
# ---------------------------------------------------------------------------

def _get(obj: Any, name: str, default: str = "") -> str:
    """Best-effort attribute/key read that never raises — `obj` may be a
    Complaint ORM row, a plain dict, or None (complaint not found / not
    linked yet)."""
    if obj is None:
        return default
    if isinstance(obj, dict):
        value = obj.get(name, default)
    else:
        value = getattr(obj, name, default)
    return "" if value is None else str(value)


def _shade_cell(cell, fill_hex: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): fill_hex,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, lines: Iterable[tuple[str, dict]]):
    cell.text = ""
    first = True
    for text, kwargs in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(text)
        run.font.size = kwargs.get("size", LABEL_SIZE)
        run.bold = kwargs.get("bold", False)
        run.italic = kwargs.get("italic", False)
        color = kwargs.get("color")
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _label_sub_blank(cell, label: str, sub: Optional[str], value: str):
    """A field cell: bold label, italic gujarati/sub label, then the filled
    value underlined (or a blank underline if there's no value)."""
    lines = [(label, {"bold": True, "size": LABEL_SIZE})]
    if sub:
        lines.append((sub, {"italic": True, "size": SUBLABEL_SIZE, "color": "555555"}))
    _set_cell_text(cell, lines)
    p = cell.add_paragraph()
    run = p.add_run(value if value else "\u00A0")
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    _underline_paragraph(p)


def _underline_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "000000",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header_row(table, num: int, title: str, sub: Optional[str] = None, span: int = 1):
    row = table.add_row()
    cell = row.cells[0]
    if span > 1:
        cell = cell.merge(row.cells[span - 1])
    _shade_cell(cell, HEADER_FILL)
    _set_cell_text(cell, [
        (f"{num}.  {title}", {"bold": True, "size": Pt(11)}),
        *([(sub, {"italic": True, "size": SUBLABEL_SIZE, "color": "555555"})] if sub else []),
    ])
    return row


def _blank_body_row(table, text: str = "", span: int = 1, lines: int = 2):
    row = table.add_row()
    cell = row.cells[0]
    if span > 1:
        cell = cell.merge(row.cells[span - 1])
    p = cell.paragraphs[0]
    run = p.add_run(text if text else "\u00A0")
    run.font.size = Pt(10)
    for _ in range(max(0, lines - 1)):
        cell.add_paragraph().add_run("\u00A0")
    return row


def _set_col_widths(table, widths_in: list[float]):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)


# ---------------------------------------------------------------------------
# main entry point
# ---------------------------------------------------------------------------

def build_fir_docx(draft: Any, complaint: Optional[Any] = None) -> io.BytesIO:
    """Build a filled-in FIR .docx for a saved `FirDraft` row.

    `draft` needs: id, complaint_id, crime_category, summary, draft_content,
    status, created_at (SQLAlchemy model instance or dict-like).
    `complaint` is optional — pass the linked Complaint row if you have one
    loaded, so complainant / place-of-occurrence fields can be filled in.
    Any attribute not present on `complaint` is simply left blank.
    """
    content = _get_dict(draft, "draft_content") or {}
    sections: list[dict] = content.get("selected_sections", []) or []
    judgments: list[dict] = content.get("selected_judgments", []) or []

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(0.7)
    section.top_margin = section.bottom_margin = Inches(0.6)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("FIRST INFORMATION REPORT")
    r.bold = True
    r.underline = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("પ્રથમ માહિતી અહેવાલ")
    sr.font.size = Pt(12)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2r = sub2.add_run("(Under Section 154 Cr.P.C.)")
    s2r.italic = True
    s2r.font.size = Pt(10)
    doc.add_paragraph()

    created_at = _get(draft, "created_at")
    created_date = created_at.split("T")[0].split(" ")[0] if created_at else ""
    draft_id = _get(draft, "id")

    # ---- Item 1: District / PS / Year / FIR No / Date ----
    t1 = doc.add_table(rows=0, cols=5)
    t1.style = "Table Grid"
    row = t1.add_row()
    for c in row.cells[1:]:
        row.cells[0].merge(c)
    _shade_cell(row.cells[0], HEADER_FILL)
    _set_cell_text(row.cells[0], [("1.  Police Station Details", {"bold": True, "size": Pt(11)})])

    row = t1.add_row()
    _label_sub_blank(row.cells[0], "District", "(જિલ્લો)", _get(complaint, "district"))
    _label_sub_blank(row.cells[1], "Police Station", "(પોલીસ સ્ટેશન)", _get(complaint, "police_station"))
    _label_sub_blank(row.cells[2], "Year", "(વર્ષ)", created_date[:4] if created_date else "")
    _label_sub_blank(row.cells[3], "FIR Draft No.", "(પ્ર.મા.અ.ક્રમાંક)", draft_id)
    _label_sub_blank(row.cells[4], "Date", "(તારીખ)", created_date)
    _set_col_widths(t1, [1.4, 1.4, 0.9, 1.6, 1.6])
    doc.add_paragraph()

    # ---- Item 2: Act & Sections (filled from selected_sections) ----
    t2 = doc.add_table(rows=0, cols=1)
    t2.style = "Table Grid"
    _section_header_row(t2, 2, "Act & Sections", "(અધિનિયમ અને કલમો)")
    if sections:
        by_act: dict[str, list[str]] = {}
        for s in sections:
            act = s.get("act_code", "—")
            label = f"Sec {s.get('section_number', '')} — {s.get('title', '')}".strip(" —")
            by_act.setdefault(act, []).append(label)
        lines = "\n".join(f"{act}: " + "; ".join(labels) for act, labels in by_act.items())
        _blank_body_row(t2, lines, lines=len(by_act) + 1)
    else:
        _blank_body_row(t2, "", lines=3)
    doc.add_paragraph()

    # ---- Item 3: Occurrence of offence ----
    t3 = doc.add_table(rows=0, cols=1)
    t3.style = "Table Grid"
    _section_header_row(t3, 3, "Occurrence of Offence", "(ગુન્હો બન્યાનો સમયગાળો)")
    occ_lines = []
    for label, key in (("Date/time of occurrence", "occurrence_date"), ("Date/time information received", "created_at")):
        val = _get(complaint, key) or (created_at if key == "created_at" else "")
        occ_lines.append(f"{label}: {val or '_' * 20}")
    _blank_body_row(t3, "\n".join(occ_lines), lines=len(occ_lines) + 1)
    doc.add_paragraph()

    # ---- Item 4: Type of information ----
    t4 = doc.add_table(rows=0, cols=1)
    t4.style = "Table Grid"
    _section_header_row(t4, 4, "Type of Information", "(માહિતીનો પ્રકાર)")
    _blank_body_row(t4, _get(complaint, "source"), lines=2)
    doc.add_paragraph()

    # ---- Item 5: Place of occurrence ----
    t5 = doc.add_table(rows=0, cols=1)
    t5.style = "Table Grid"
    _section_header_row(t5, 5, "Place of Occurrence", "(ઘટનાનું સ્થળ)")
    _blank_body_row(t5, _get(complaint, "address") or _get(complaint, "location"), lines=3)
    doc.add_paragraph()

    # ---- Item 6: Complainant / Informant ----
    t6 = doc.add_table(rows=0, cols=2)
    t6.style = "Table Grid"
    row = t6.add_row()
    row.cells[0].merge(row.cells[1])
    _shade_cell(row.cells[0], HEADER_FILL)
    _set_cell_text(row.cells[0], [
        ("6.  Complainant / Informant", {"bold": True, "size": Pt(11)}),
        ("(ફરિયાદી / બાતમીદાર)", {"italic": True, "size": SUBLABEL_SIZE, "color": "555555"}),
    ])
    row = t6.add_row()
    _label_sub_blank(row.cells[0], "(a) Name", "(નામ)", _get(complaint, "complainant_name"))
    _label_sub_blank(row.cells[1], "(b) Father's / Husband's Name", "(પિતા/પતિનું નામ)", _get(complaint, "father_or_husband_name"))
    row = t6.add_row()
    _label_sub_blank(row.cells[0], "(e) Occupation", "(ધંધો)", _get(complaint, "complainant_occupation"))
    _label_sub_blank(row.cells[1], "(g) Address", "(સરનામું)", _get(complaint, "complainant_address"))
    _set_col_widths(t6, [3.55, 3.55])
    doc.add_paragraph()

    # ---- Item 7: Accused table ----
    accused = content.get("accused") or (complaint.get("accused") if isinstance(complaint, dict) else getattr(complaint, "accused", None)) or []
    t7h = doc.add_table(rows=0, cols=1)
    t7h.style = "Table Grid"
    _section_header_row(t7h, 7, "Details of Known / Suspected / Unknown Accused",
                         "(ઓળખાયેલ/શકમંદ/વણઓળખાયેલ આરોપીની વિગતો)")
    t7 = doc.add_table(rows=1, cols=3)
    t7.style = "Table Grid"
    hdr = t7.rows[0]
    for cell, text in zip(hdr.cells, ["Accused Name (તહોમતદારનું નામ)", "Age (ઉંમર)", "Address (સરનામું)"]):
        _shade_cell(cell, "F2F2F2")
        _set_cell_text(cell, [(text, {"bold": True, "size": Pt(9)})])
    rows_needed = max(4, len(accused))
    for i in range(rows_needed):
        row = t7.add_row()
        a = accused[i] if i < len(accused) else {}
        row.cells[0].text = str(a.get("name", "")) if isinstance(a, dict) else ""
        row.cells[1].text = str(a.get("age", "")) if isinstance(a, dict) else ""
        row.cells[2].text = str(a.get("address", "")) if isinstance(a, dict) else ""
    _set_col_widths(t7, [3.0, 1.3, 2.8])
    doc.add_paragraph()

    # ---- Item 12: First Information contents (case summary + citations) ----
    t12 = doc.add_table(rows=0, cols=1)
    t12.style = "Table Grid"
    _section_header_row(t12, 12, "First Information Contents", "(પ્રથમ માહિતી અહેવાલના વિગતો)")
    narrative = _get(draft, "summary") or ""
    body_row = t12.add_row()
    cell = body_row.cells[0]
    p = cell.paragraphs[0]
    p.add_run("Complaint (ફરિયાદ):").bold = True
    body_p = cell.add_paragraph()
    body_p.add_run(narrative if narrative else "\u00A0" * 10)

    if judgments:
        cite_p = cell.add_paragraph()
        cite_p.paragraph_format.space_before = Pt(10)
        cite_p.add_run("Cited landmark judgments:").bold = True
        for j in judgments:
            jp = cell.add_paragraph(style="List Bullet")
            title = j.get("case_title", "")
            court = j.get("court", "")
            date = j.get("case_date", "")
            outcome = j.get("bail_outcome", "")
            jp.add_run(f"{title} — {court}, {date}" + (f" ({outcome})" if outcome else ""))
    for _ in range(3):
        cell.add_paragraph().add_run("\u00A0")
    doc.add_paragraph()

    # ---- Item 13: Action taken ----
    t13 = doc.add_table(rows=0, cols=1)
    t13.style = "Table Grid"
    _section_header_row(t13, 13, "Action Taken", "(લીધેલ પગલાં)")
    status = _get(draft, "status") or "draft"
    crime_category = _get(draft, "crime_category")
    action_row = t13.add_row()
    c = action_row.cells[0]
    c.paragraphs[0].add_run(f"Crime category: {crime_category or '_' * 20}")
    c.add_paragraph().add_run(f"Draft status: {status}")
    c.add_paragraph().add_run("\u2610  (1) Registered the case and took up the investigation, or")
    c.add_paragraph().add_run("\u2610  (2) Directed (Name of I.O.) to take up the investigation: " + "_" * 30)
    doc.add_paragraph()

    # ---- Item 14/15: Signatures ----
    t14 = doc.add_table(rows=0, cols=2)
    t14.style = "Table Grid"
    row = t14.add_row()
    _label_sub_blank(row.cells[0], "14. Signature / Thumb Impression of Complainant", "(ફરિયાદી/બાતમીદારની સહી)", "")
    _label_sub_blank(row.cells[1], "Signature of Officer in Charge", "(અધિકારીની સહી)", _get(draft, "approved_by"))
    _set_col_widths(t14, [3.55, 3.55])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _get_dict(obj: Any, name: str) -> Optional[dict]:
    if obj is None:
        return None
    value = obj.get(name) if isinstance(obj, dict) else getattr(obj, name, None)
    return value