import os
from docx import Document

templates_dir = r"c:\Users\vyomi\OneDrive\Desktop\CaseCraftAI\backend\app\services\templates"
templates = [
    "purvani_chargesheet.docx",
    "medical_treatment_letter.docx",
    "remand_request_letter.docx",
    "seizure_receipt.docx",
    "court_custody_request.docx",
    "accused_panchanama.docx",
    "accused_face_identification_form.docx"
]

for t in templates:
    filepath = os.path.join(templates_dir, t)
    if not os.path.exists(filepath):
        print(f"File not found: {t}")
        continue
    doc = Document(filepath)
    print(f"\n==================== TEMPLATE: {t} ====================")
    
    # Print paragraphs that contain any brackets
    for i, p in enumerate(doc.paragraphs):
        if any(c in p.text for c in ['{', '}', '[', ']']):
            print(f"P{i}: {p.text.strip()}")
            
    # Print table cells containing any brackets
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if any(c in cell.text for c in ['{', '}', '[', ']']):
                    print(f"T{t_idx}R{r_idx}C{c_idx}: {cell.text.strip()}")
